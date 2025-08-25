#!/usr/bin/env python3
"""
Enhanced FastAPI Backend with RAG for Document Question Answering
Uses FAISS for vector search, Gemini for generation, and Supabase for storage
Production-ready with intelligent chunking and semantic retrieval - NO LOCAL STORAGE
"""

from fastapi import FastAPI, HTTPException, Depends, status, Request, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, field_validator
from typing import List, Optional, Dict, Any
import httpx
import asyncio
import logging
import os
import sys  # Add this line
from dotenv import load_dotenv
import time
import hashlib
import json
import uuid
from pathlib import Path
import pickle
import io
from io import BytesIO

from typing import Tuple  # Add this to existing typing imports

import signal

# Document processing imports
import PyPDF2
import docx
import email

# RAG and embedding imports
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
import nltk
from nltk.tokenize import sent_tokenize
import re
from urllib.parse import urlparse, parse_qs
from pathlib import Path
import mimetypes

# LLM integration (Google Generative AI)
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold

# Supabase storage utilities - FIXED IMPORTS
from supabase_utils import (
    upload_file_to_supabase,
    download_file_from_supabase,
    download_document_content,
    get_public_url,
    delete_file_from_supabase,
    get_supabase_manager,
    list_supabase_files,
    test_supabase_upload_standalone
)

import tempfile

# Configure logging with more detailed format for better debugging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    try:
        nltk.download('punkt', quiet=True)
    except:
        try:
            nltk.download('punkt_tab', quiet=True)
        except:
            logger.warning("Could not download NLTK punkt tokenizer")
            pass

# Add this class after the imports and before the existing code
class GracefulKiller:
    """Handle graceful shutdown for Cloud Run"""
    kill_now = False
    def __init__(self):
        signal.signal(signal.SIGINT, self._exit_gracefully)
        signal.signal(signal.SIGTERM, self._exit_gracefully)

    def _exit_gracefully(self, signum, frame):
        logger.info(f"Received signal {signum}, shutting down gracefully...")
        self.kill_now = True

# Add this line after the GracefulKiller class definition
killer = GracefulKiller()

# Load environment variables FIRST
load_dotenv()

# Initialize FastAPI app
app = FastAPI(
    title="RAG-Powered Document QA API with Full Supabase Storage",
    description="Advanced Document Question Answering with Retrieval-Augmented Generation and Complete Cloud Storage",
    version="2.2.0",
    docs_url="/api/v1/docs",
    redoc_url="/api/v1/redoc"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security
security = HTTPBearer()
VALID_TOKEN = os.getenv("VALID_TOKEN")

# Configuration
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"  # Lightweight but effective
CHUNK_SIZE = 1500  # Characters per chunk
CHUNK_OVERLAP = 200  # Overlap between chunks
TOP_K_RETRIEVAL = 8  # Number of chunks to retrieve
MAX_CONTEXT_LENGTH = 10000  # Max context for Gemini

# DEBUG: Log all critical environment variables on startup
logger.info("🔧 ENVIRONMENT VARIABLES DEBUG:")
logger.info(f"   SUPABASE_URL: {os.getenv('SUPABASE_URL')}")
logger.info(
    f"   SUPABASE_KEY: {'*' * (len(os.getenv('SUPABASE_KEY', '')) - 8) + os.getenv('SUPABASE_KEY', '')[-8:] if os.getenv('SUPABASE_KEY') else 'NOT_SET'}")
logger.info(f"   SUPABASE_BUCKET: {os.getenv('SUPABASE_BUCKET', 'documents')}")
logger.info(f"   GEMINI_API_KEY: {'SET' if GEMINI_API_KEY else 'NOT_SET'}")
logger.info(f"   VALID_TOKEN: {'SET' if VALID_TOKEN else 'NOT_SET'}")

# Validate they exist
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY environment variable is required")
if not VALID_TOKEN:
    raise ValueError("VALID_TOKEN environment variable is required")

# Validate Supabase credentials
if not os.getenv("SUPABASE_URL"):
    raise ValueError("SUPABASE_URL environment variable is required")
if not os.getenv("SUPABASE_KEY"):
    raise ValueError("SUPABASE_KEY environment variable is required")

# Configure Gemini
try:
    genai.configure(api_key=GEMINI_API_KEY)
    logger.info("✅ Gemini API configured successfully")
except Exception as e:
    logger.error(f"❌ Failed to configure Gemini API: {e}")


# Request/Response Models
class DocumentQARequest(BaseModel):
    documents: Optional[str] = None  # URL or Supabase file path
    questions: List[str]
    document_id: Optional[str] = None  # For pre-processed documents

    @field_validator('questions')
    @classmethod
    def validate_questions(cls, v):
        if not v or len(v) == 0:
            raise ValueError("At least one question is required")
        if len(v) > 10:
            raise ValueError("Maximum 10 questions allowed per request")
        return v


class GlobalQueryRequest(BaseModel):
    query: str
    top_k: int = 10
    max_docs: int = 5
    document_ids: Optional[List[str]] = None  # NEW: Filter by document IDs or filenames

    @field_validator('query')
    @classmethod
    def validate_query(cls, v):
        if not v or not v.strip():
            raise ValueError("Query cannot be empty")
        if len(v) > 500:
            raise ValueError("Query too long (max 500 characters)")
        return v.strip()

    @field_validator('document_ids')
    @classmethod
    def validate_document_ids(cls, v):
        if v is not None:
            if len(v) == 0:
                raise ValueError("document_ids cannot be empty if provided")
            if len(v) > 20:
                raise ValueError("Maximum 20 documents allowed per query")
        return v



class DocumentUploadResponse(BaseModel):
    document_id: str
    filename: str
    status: str
    chunks_created: int
    message: str
    supabase_path: str
    public_url: Optional[str] = None


class DocumentQAResponse(BaseModel):
    answers: List[Dict[str, Any]]
    document_id: str
    retrieval_info: Optional[Dict[str, Any]] = None


class ErrorResponse(BaseModel):
    error: str
    details: Optional[str] = None


# Authentication dependency
async def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    if credentials.credentials != VALID_TOKEN:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid authentication token",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return credentials.credentials


class DocumentProcessor:
    """Enhanced document processor with intelligent chunking and Supabase integration"""

    @staticmethod
    def extract_text_from_pdf(content: bytes) -> str:
        """Extract text from PDF with better error handling"""
        try:
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""

            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    logger.warning(f"Error processing page {page_num + 1}: {e}")
                    continue

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to process PDF: {str(e)}"
            )

    @staticmethod
    def extract_text_from_docx(content: bytes) -> str:
        """Extract text from DOCX with enhanced processing"""
        try:
            docx_file = BytesIO(content)
            doc = docx.Document(docx_file)
            text = ""

            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to process DOCX: {str(e)}"
            )

    @staticmethod
    def extract_text_from_email(content: bytes) -> str:
        """Extract text from email content"""
        try:
            email_str = content.decode('utf-8', errors='ignore')
            msg = email.message_from_string(email_str)

            text = ""

            # Extract headers
            for header in ['From', 'To', 'Subject', 'Date']:
                if msg.get(header):
                    text += f"{header}: {msg.get(header)}\n"
            text += "\n"

            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            text += payload.decode('utf-8', errors='ignore') + "\n"
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    text += payload.decode('utf-8', errors='ignore')

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting email text: {e}")
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to process email: {str(e)}"
            )

    @classmethod
    def intelligent_chunking(cls, text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[
        Dict[str, Any]]:
        """
        Intelligent text chunking that preserves semantic meaning
        """
        # Clean the text
        text = cls._clean_text(text)

        # Split into sentences
        sentences = sent_tokenize(text)

        chunks = []
        current_chunk = ""
        current_chunk_sentences = []

        for i, sentence in enumerate(sentences):
            # Check if adding this sentence would exceed chunk size
            potential_chunk = current_chunk + " " + sentence if current_chunk else sentence

            if len(potential_chunk) <= chunk_size:
                current_chunk = potential_chunk
                current_chunk_sentences.append(sentence)
            else:
                # Save current chunk if it has content
                if current_chunk:
                    chunks.append({
                        "text": current_chunk.strip(),
                        "chunk_id": len(chunks),
                        "sentence_count": len(current_chunk_sentences),
                        "char_count": len(current_chunk)
                    })

                # Start new chunk with overlap
                if overlap > 0 and current_chunk_sentences:
                    # Calculate how many sentences to include for overlap
                    overlap_sentences = []
                    overlap_chars = 0

                    for sent in reversed(current_chunk_sentences):
                        if overlap_chars + len(sent) <= overlap:
                            overlap_sentences.insert(0, sent)
                            overlap_chars += len(sent)
                        else:
                            break

                    current_chunk = " ".join(overlap_sentences + [sentence])
                    current_chunk_sentences = overlap_sentences + [sentence]
                else:
                    current_chunk = sentence
                    current_chunk_sentences = [sentence]

        # Add the last chunk
        if current_chunk:
            chunks.append({
                "text": current_chunk.strip(),
                "chunk_id": len(chunks),
                "sentence_count": len(current_chunk_sentences),
                "char_count": len(current_chunk)
            })

        logger.info(f"Created {len(chunks)} intelligent chunks")
        return chunks

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize text"""
        # Handle None or empty text
        if not text:
            return ""

        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove page breaks and form feeds
        text = re.sub(r'[\f\r]+', '\n', text)

        # Normalize quotes - using Unicode escapes to avoid syntax issues
        text = re.sub(r'["""]', '"', text)  # Double quotes
        text = re.sub(r'[\u2018\u2019]', "'", text)  # Single quotes (left and right)

        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)

        return text.strip()

    @classmethod
    async def process_document_from_source(cls, source: str, source_type: str = "auto") -> tuple[str, str, str, str]:
        """Process document from URL or Supabase storage and return text, document ID, original filename, and sanitized filename"""
        # Generate document ID based on source
        doc_id = hashlib.md5(source.encode()).hexdigest()

        # Determine source type
        if source_type == "auto":
            if GoogleDriveProcessor.is_google_drive_link(source):
                source_type = "google_drive"
            elif source.startswith(('http://', 'https://')):
                source_type = "url"
            else:
                source_type = "supabase_path"

        # Handle Google Drive links
        if source_type == "google_drive":
            # Get original filename BEFORE converting to direct download URL
            original_filename, sanitized_filename = await cls._extract_filename_from_url(source)
            # Convert to direct download URL
            direct_url = GoogleDriveProcessor.convert_to_direct_download(source)
            source = direct_url
        else:
            # Extract filename from regular URL or use default
            original_filename, sanitized_filename = await cls._extract_filename_from_url(source)

        # Download document content
        content, _ = await download_document_content(source)

        logger.info(
            f"Processing document from {source_type}: {original_filename} (sanitized: {sanitized_filename}) ({len(content)} bytes)")

        # Determine file type and extract text
        text = await cls._extract_text_by_content_type(content, source, original_filename)

        if not text.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Document appears to be empty or contains no extractable text"
            )

        return text, doc_id, original_filename, sanitized_filename

    @staticmethod
    async def _extract_filename_from_url(url: str) -> Tuple[str, str]:
        """Extract filename from URL - returns (original_filename, sanitized_filename)"""
        try:
            if 'drive.google.com' in url:
                # Extract file ID from Google Drive URL
                file_id = None

                # Try different patterns
                patterns = [
                    r'/file/d/([a-zA-Z0-9-_]+)',
                    r'id=([a-zA-Z0-9-_]+)',
                    r'/d/([a-zA-Z0-9-_]+)'
                ]

                for pattern in patterns:
                    match = re.search(pattern, url)
                    if match:
                        file_id = match.group(1)
                        break

                if file_id:
                    # Get real filename from Drive API with improved logic
                    metadata = await GoogleDriveProcessor.get_drive_file_metadata(file_id)
                    original_name = metadata["name"]

                    # Only use fallback if the returned name looks like our fallback pattern
                    if not original_name.startswith('drive_document_') and not original_name.startswith('gdrive_doc_'):
                        sanitized_name = GoogleDriveProcessor.sanitize_filename(original_name)
                        logger.info(f"✅ Using real Drive filename: {original_name}")
                        return original_name, sanitized_name
                    else:
                        logger.warning(f"Drive API returned fallback name: {original_name}")
                        # Try to extract from URL as backup
                        parsed = urlparse(url)
                        params = parse_qs(parsed.query)

                        for param_name in ['filename', 'name', 'title']:
                            if param_name in params and params[param_name][0]:
                                url_filename = params[param_name][0]
                                if '.' in url_filename and len(url_filename) > 3:
                                    sanitized = GoogleDriveProcessor.sanitize_filename(url_filename)
                                    logger.info(f"✅ Using filename from URL params: {url_filename}")
                                    return url_filename, sanitized

                        # Final fallback with better naming
                        timestamp = int(time.time())
                        final_fallback = f"google_drive_document_{timestamp}.pdf"
                        return final_fallback, final_fallback

            # Rest of the method stays the same for non-Google Drive URLs...
            # [Keep existing code for regular URLs]

        except Exception as e:
            logger.warning(f"Error extracting filename from {url}: {e}")
            timestamp = int(time.time())
            default_name = f"document_{timestamp}.pdf"
            return default_name, default_name

    @classmethod
    async def _extract_text_by_content_type(cls, content: bytes, source: str, filename: str) -> str:
        """Extract text based on content type detection"""
        source_lower = source.lower()
        filename_lower = filename.lower()

        # Detect by content signature first
        if content.startswith(b'%PDF'):
            return cls.extract_text_from_pdf(content)
        elif content.startswith(b'PK'):  # ZIP-based formats like DOCX
            return cls.extract_text_from_docx(content)
        elif b'From:' in content[:1000] or b'Subject:' in content[:1000]:
            return cls.extract_text_from_email(content)

        # Fallback to file extension
        if any(ext in filename_lower for ext in ['.pdf']) or 'pdf' in source_lower:
            return cls.extract_text_from_pdf(content)
        elif any(ext in filename_lower for ext in ['.docx', '.doc']) or 'docx' in source_lower:
            return cls.extract_text_from_docx(content)
        elif '.eml' in filename_lower:
            return cls.extract_text_from_email(content)
        else:
            # Try as plain text
            try:
                return content.decode('utf-8', errors='ignore')
            except Exception:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Unsupported document format. Supported: PDF, DOCX, EML, TXT"
                )
    @classmethod
    async def process_uploaded_file(cls, content: bytes, filename: str) -> str:
        """Process uploaded file content and return extracted text"""
        logger.info(f"Processing uploaded file: {filename} ({len(content)} bytes)")

        # Determine file type and extract text
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf') or content.startswith(b'%PDF'):
            text = cls.extract_text_from_pdf(content)
        elif filename_lower.endswith('.docx') or content.startswith(b'PK'):
            text = cls.extract_text_from_docx(content)
        elif filename_lower.endswith('.eml'):
            text = cls.extract_text_from_email(content)
        else:
            try:
                text = content.decode('utf-8', errors='ignore')
            except Exception:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Unsupported document format. Supported: PDF, DOCX, EML, TXT"
                )

        if not text.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Document appears to be empty or contains no extractable text"
            )

        return text

    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """Sanitize filename for safe storage while preserving readability"""
        # Remove or replace problematic characters
        sanitized = filename.replace(" ", "_").replace("/", "_").replace("\\", "_")
        sanitized = re.sub(r'[<>:"|?*]', '_', sanitized)
        # Remove multiple underscores
        sanitized = re.sub(r'_+', '_', sanitized)
        return sanitized.strip('_')


class GoogleDriveProcessor:
    """Handle Google Drive folder and file processing"""

    @staticmethod
    def is_google_drive_link(url: str) -> bool:
        """Check if URL is a Google Drive link"""
        return 'drive.google.com' in url or 'docs.google.com' in url

    @staticmethod
    def extract_folder_id(url: str) -> str:
        """Extract folder ID from Google Drive URL"""
        patterns = [
            r'/folders/([a-zA-Z0-9-_]+)',
            r'id=([a-zA-Z0-9-_]+)',
            r'/d/([a-zA-Z0-9-_]+)'
        ]

        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)

        raise ValueError("Could not extract folder ID from Google Drive URL")

    @staticmethod
    async def get_drive_folder_files(folder_url: str) -> List[Dict[str, str]]:
        """Get list of files from Google Drive folder - Fixed version with better error handling"""
        try:
            folder_id = GoogleDriveProcessor.extract_folder_id(folder_url)
            logger.info(f"Processing Google Drive folder: {folder_id}")

            # Try to make the folder publicly accessible by converting to a viewable URL
            public_folder_url = f"https://drive.google.com/drive/folders/{folder_id}"

            # Use HTML parsing approach since API requires authentication
            timeout = httpx.Timeout(30.0, connect=10.0)
            async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
                # Try to access the folder page
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = await client.get(public_folder_url, headers=headers)

                if response.status_code != 200:
                    logger.error(f"Cannot access Google Drive folder. Status: {response.status_code}")
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"Cannot access Google Drive folder. Make sure the folder is publicly accessible. Status: {response.status_code}"
                    )

                content = response.text

                # Extract file IDs and names from HTML
                files = []

                # Look for file patterns in the HTML
                import re

                # Pattern for file data in Drive's HTML
                file_patterns = [
                    r'"(\w{28,})".*?"([^"]*\.(?:pdf|docx?|txt|eml))"',  # File ID and name
                    r'/file/d/([a-zA-Z0-9-_]{28,})',  # Just file IDs
                ]

                found_files = set()  # Avoid duplicates

                # Inside the folder processing method, update the file detection section:
                for pattern in file_patterns:
                    matches = re.finditer(pattern, content, re.IGNORECASE)
                    for match in matches:
                        if len(match.groups()) == 2:
                            file_id, detected_filename = match.groups()
                        else:
                            file_id = match.group(1)
                            detected_filename = None

                        if file_id not in found_files and len(file_id) >= 28:
                            found_files.add(file_id)

                            # Get real filename using our improved metadata function
                            try:
                                metadata = await GoogleDriveProcessor.get_drive_file_metadata(file_id)
                                real_filename = metadata["name"]
                                logger.info(f"📄 Got real filename for {file_id[:8]}: {real_filename}")
                            except Exception as meta_error:
                                logger.warning(f"Failed to get metadata for {file_id}: {meta_error}")
                                real_filename = detected_filename or f"drive_document_{file_id[:8]}.pdf"

                            # Create download URL
                            download_url = f"https://drive.google.com/uc?export=download&id={file_id}&confirm=t"

                            files.append({
                                "file_id": file_id,
                                "download_url": download_url,
                                "filename": real_filename,  # Use the real filename here
                                "original_filename": real_filename
                            })

                if not files:
                    logger.warning("No files found in Google Drive folder")
                    # Try alternative extraction method
                    file_id_matches = re.findall(r'/file/d/([a-zA-Z0-9-_]{25,})', content)
                    for file_id in list(set(file_id_matches))[:10]:  # Limit to 10 files
                        download_url = f"https://drive.google.com/uc?export=download&id={file_id}&confirm=t"
                        files.append({
                            "file_id": file_id,
                            "download_url": download_url,
                            "filename": f"document_{file_id[:8]}.pdf",
                            "original_filename": f"document_{file_id[:8]}.pdf"
                        })

                logger.info(f"Found {len(files)} files in Google Drive folder")
                return files

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error processing Google Drive folder: {e}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to process Google Drive folder: {str(e)}. Make sure the folder is publicly accessible."
            )

    @staticmethod
    async def _fallback_folder_parsing(folder_url: str) -> List[Dict[str, str]]:
        """Fallback method using HTML parsing"""
        try:
            timeout = httpx.Timeout(30.0, connect=10.0)
            async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
                response = await client.get(folder_url)
                response.raise_for_status()
                content = response.text

                files = []
                for match in re.finditer(r'/file/d/([a-zA-Z0-9-_]+)', content):
                    file_id = match.group(1)
                    download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
                    files.append({
                        "file_id": file_id,
                        "download_url": download_url,
                        "filename": f"document_{file_id}.pdf",  # Default fallback
                        "original_filename": f"document_{file_id}.pdf"
                    })

                return files
        except Exception as e:
            logger.error(f"Fallback parsing failed: {e}")
            return []

    @staticmethod
    def convert_to_direct_download(url: str) -> str:
        """Convert Google Drive share URL to direct download URL"""
        if '/file/d/' in url:
            file_id_match = re.search(r'/file/d/([a-zA-Z0-9-_]+)', url)
            if file_id_match:
                file_id = file_id_match.group(1)
                return f"https://drive.google.com/uc?export=download&id={file_id}"

        if 'id=' in url:
            parsed = urlparse(url)
            params = parse_qs(parsed.query)
            if 'id' in params:
                file_id = params['id'][0]
                return f"https://drive.google.com/uc?export=download&id={file_id}"

        return url

    @staticmethod
    async def get_drive_file_metadata(file_id: str) -> Dict[str, str]:
        """Get file metadata from Google Drive using multiple approaches with better error handling"""
        try:
            # Method 1: Try Google Drive API v3 (most reliable)
            api_url = f"https://www.googleapis.com/drive/v3/files/{file_id}?fields=name,mimeType,size"

            timeout = httpx.Timeout(15.0, connect=5.0)
            async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
                try:
                    response = await client.get(api_url)
                    if response.status_code == 200:
                        metadata = response.json()
                        original_name = metadata.get("name", "").strip()

                        # Validate the name is meaningful
                        if original_name and len(original_name) > 0 and not original_name.startswith('Untitled'):
                            logger.info(f"✅ Got real filename from Drive API: {original_name}")
                            return {
                                "name": original_name,
                                "mime_type": metadata.get("mimeType", "application/pdf"),
                                "size": metadata.get("size", "unknown")
                            }
                    else:
                        logger.warning(f"Drive API returned status {response.status_code}")
                except Exception as api_error:
                    logger.warning(f"Drive API v3 failed: {api_error}")

            # Method 2: Try alternative API endpoint
            alt_api_url = f"https://www.googleapis.com/drive/v2/files/{file_id}"
            try:
                response = await client.get(alt_api_url)
                if response.status_code == 200:
                    metadata = response.json()
                    original_name = metadata.get("title", "").strip()

                    if original_name and len(original_name) > 0 and not original_name.startswith('Untitled'):
                        logger.info(f"✅ Got filename from Drive API v2: {original_name}")
                        return {
                            "name": original_name,
                            "mime_type": metadata.get("mimeType", "application/pdf"),
                            "size": str(metadata.get("fileSize", "unknown"))
                        }
            except Exception as api2_error:
                logger.warning(f"Drive API v2 failed: {api2_error}")

            # Method 3: Try HTML scraping approach
            try:
                file_url = f"https://drive.google.com/file/d/{file_id}/view"
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }

                response = await client.get(file_url, headers=headers)
                if response.status_code == 200:
                    content = response.text

                    # Look for filename in HTML title tag
                    import re
                    title_match = re.search(r'<title>([^<]+)</title>', content, re.IGNORECASE)
                    if title_match:
                        title = title_match.group(1).strip()
                        # Remove " - Google Drive" suffix
                        if title.endswith(' - Google Drive'):
                            title = title[:-13].strip()

                        if title and len(title) > 3 and not title.startswith('Google Drive'):
                            logger.info(f"✅ Got filename from HTML title: {title}")
                            # Add extension if missing
                            if '.' not in title:
                                title += '.pdf'
                            return {
                                "name": title,
                                "mime_type": "application/pdf",
                                "size": "unknown"
                            }

                    # Look for filename in meta tags
                    meta_patterns = [
                        r'<meta property="og:title" content="([^"]+)"',
                        r'<meta name="title" content="([^"]+)"',
                        r'"name":"([^"]+\.(?:pdf|docx?|txt|eml))"'
                    ]

                    for pattern in meta_patterns:
                        matches = re.finditer(pattern, content, re.IGNORECASE)
                        for match in matches:
                            filename = match.group(1).strip()
                            if filename and len(filename) > 3:
                                logger.info(f"✅ Got filename from HTML meta: {filename}")
                                return {
                                    "name": filename,
                                    "mime_type": "application/pdf",
                                    "size": "unknown"
                                }
            except Exception as html_error:
                logger.warning(f"HTML scraping failed: {html_error}")

            # All methods failed - return descriptive fallback
            logger.warning(f"All filename extraction methods failed for file {file_id}")
            fallback_name = f"drive_document_{file_id[:8]}.pdf"

            return {
                "name": fallback_name,
                "mime_type": "application/pdf",
                "size": "unknown"
            }

        except Exception as e:
            logger.error(f"Complete metadata extraction failed for {file_id}: {e}")
            return {
                "name": f"drive_document_{file_id[:8]}.pdf",
                "mime_type": "application/pdf",
                "size": "unknown"
            }

    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """Sanitize filename for safe storage while preserving readability"""
        # Remove or replace problematic characters
        sanitized = filename.replace(" ", "_").replace("/", "_").replace("\\", "_")
        sanitized = re.sub(r'[<>:"|?*]', '_', sanitized)
        # Remove multiple underscores
        sanitized = re.sub(r'_+', '_', sanitized)
        return sanitized.strip('_')

    @staticmethod
    async def get_drive_folder_files_recursive(folder_url: str, current_path: str = "", max_depth: int = 5,
                                               current_depth: int = 0) -> List[Dict[str, str]]:
        """
        Recursively get all files from Google Drive folder and its subfolders

        Args:
            folder_url (str): Google Drive folder URL
            current_path (str): Current folder path for nested folders
            max_depth (int): Maximum recursion depth to prevent infinite loops
            current_depth (int): Current recursion depth

        Returns:
            List[Dict[str, str]]: List of all files with folder path information
        """
        if current_depth >= max_depth:
            logger.warning(f"Maximum recursion depth ({max_depth}) reached for path: {current_path}")
            return []

        try:
            folder_id = GoogleDriveProcessor.extract_folder_id(folder_url)
            logger.info(f"📁 Processing Google Drive folder (depth {current_depth}): {current_path or 'root'}")

            # Get folder contents using enhanced HTML parsing
            all_items = await GoogleDriveProcessor._get_folder_contents_enhanced(folder_id)

            all_files = []

            for item in all_items:
                if item.get("type") == "file":
                    # It's a file - add folder path information
                    file_info = {
                        "file_id": item["file_id"],
                        "download_url": item["download_url"],
                        "filename": item["filename"],
                        "original_filename": item["original_filename"],
                        "folder_path": current_path
                    }
                    all_files.append(file_info)

                elif item.get("type") == "folder":
                    # It's a subfolder - recursively process it
                    subfolder_name = item.get("folder_name", "unknown_folder")
                    subfolder_path = f"{current_path}/{subfolder_name}" if current_path else subfolder_name

                    logger.info(f"📂 Found subfolder: {subfolder_path}")

                    try:
                        # Construct subfolder URL
                        subfolder_url = f"https://drive.google.com/drive/folders/{item['folder_id']}"

                        # Recursively process subfolder
                        subfolder_files = await GoogleDriveProcessor.get_drive_folder_files_recursive(
                            subfolder_url,
                            subfolder_path,
                            max_depth,
                            current_depth + 1
                        )

                        all_files.extend(subfolder_files)
                        logger.info(f"📁 Added {len(subfolder_files)} files from subfolder: {subfolder_path}")

                    except Exception as subfolder_error:
                        logger.warning(f"⚠️ Failed to process subfolder {subfolder_path}: {subfolder_error}")
                        continue

            logger.info(f"✅ Found {len(all_files)} total files in folder path: {current_path or 'root'}")
            return all_files

        except Exception as e:
            logger.error(f"❌ Error processing folder at depth {current_depth}: {e}")
            return []

    @staticmethod
    async def _get_folder_contents_enhanced(folder_id: str) -> List[Dict[str, str]]:
        """
        Enhanced folder content parsing that can detect both files and subfolders

        Returns:
            List containing both files and folders with type information
        """
        try:
            public_folder_url = f"https://drive.google.com/drive/folders/{folder_id}"

            timeout = httpx.Timeout(30.0, connect=10.0)
            async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = await client.get(public_folder_url, headers=headers)

                if response.status_code != 200:
                    logger.error(f"Cannot access Google Drive folder {folder_id}. Status: {response.status_code}")
                    return []

                content = response.text
                all_items = []

                # Enhanced regex patterns for both files and folders
                patterns = {
                    "files": [
                        r'"(\w{28,})".*?"([^"]*\.(?:pdf|docx?|txt|eml))".*?"application/',  # Files with extensions
                        r'/file/d/([a-zA-Z0-9-_]{28,})',  # File IDs
                    ],
                    "folders": [
                        r'/folders/([a-zA-Z0-9-_]{28,})"[^"]*"([^"]+)".*?"folder"',  # Folder pattern
                        r'"(\w{28,})".*?"([^"]+)".*?"folder"',  # Alternative folder pattern
                    ]
                }

                # Process files
                found_file_ids = set()
                for pattern in patterns["files"]:
                    matches = re.finditer(pattern, content, re.IGNORECASE)
                    for match in matches:
                        if len(match.groups()) >= 2:
                            file_id, filename = match.groups()[:2]
                        else:
                            file_id = match.group(1)
                            filename = f"document_{file_id[:8]}.pdf"

                        if file_id not in found_file_ids and len(file_id) >= 28:
                            found_file_ids.add(file_id)

                            download_url = f"https://drive.google.com/uc?export=download&id={file_id}&confirm=t"

                            all_items.append({
                                "type": "file",
                                "file_id": file_id,
                                "download_url": download_url,
                                "filename": filename,
                                "original_filename": filename
                            })

                # Process folders
                found_folder_ids = set()
                for pattern in patterns["folders"]:
                    matches = re.finditer(pattern, content, re.IGNORECASE)
                    for match in matches:
                        if len(match.groups()) >= 2:
                            folder_id, folder_name = match.groups()[:2]

                            if folder_id not in found_folder_ids and len(folder_id) >= 28:
                                found_folder_ids.add(folder_id)

                                all_items.append({
                                    "type": "folder",
                                    "folder_id": folder_id,
                                    "folder_name": folder_name.strip()
                                })

                # Alternative file detection if no files found with primary method
                if not found_file_ids:
                    file_id_matches = re.findall(r'/file/d/([a-zA-Z0-9-_]{25,})', content)
                    for file_id in list(set(file_id_matches))[:20]:  # Limit to 20 files
                        download_url = f"https://drive.google.com/uc?export=download&id={file_id}&confirm=t"
                        all_items.append({
                            "type": "file",
                            "file_id": file_id,
                            "download_url": download_url,
                            "filename": f"document_{file_id[:8]}.pdf",
                            "original_filename": f"document_{file_id[:8]}.pdf"
                        })

                files_count = len([item for item in all_items if item["type"] == "file"])
                folders_count = len([item for item in all_items if item["type"] == "folder"])

                logger.info(f"📊 Found {files_count} files and {folders_count} subfolders in folder {folder_id}")
                return all_items

        except Exception as e:
            logger.error(f"❌ Enhanced folder parsing failed for {folder_id}: {e}")
            return []

class VectorStore:
    """FAISS-based vector store with direct Supabase storage - NO BUFFER/CACHE"""

    def __init__(self):
        self.embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
        self.dimension = self.embedding_model.get_sentence_embedding_dimension()
        # REMOVED: No more in-memory cache/buffers
        # self.indexes = {}
        # self.chunks = {}
        logger.info(f"✅ Initialized vector store with {EMBEDDING_MODEL_NAME} (dim: {self.dimension})")
        logger.info("🚫 NO BUFFER MODE - All data stored directly in Supabase")

    async def create_embeddings(self, document_id: str, chunks: List[Dict[str, Any]],
                                file_name: str = None, source_info: Dict[str, Any] = None) -> int:
        """Create embeddings and store directly in Supabase - NO BUFFER"""
        try:
            logger.info(f"Creating embeddings for {len(chunks)} chunks - DIRECT TO SUPABASE")

            # Extract text from chunks
            texts = [chunk["text"] for chunk in chunks]

            # Generate embeddings in batches
            batch_size = 32
            all_embeddings = []

            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i + batch_size]
                batch_embeddings = self.embedding_model.encode(batch_texts, show_progress_bar=False)
                all_embeddings.append(batch_embeddings)
                logger.info(f"Processed batch {i // batch_size + 1}/{(len(texts) - 1) // batch_size + 1}")

            # Combine all embeddings
            embeddings = np.vstack(all_embeddings).astype('float32')

            # Create FAISS index
            index = faiss.IndexFlatIP(self.dimension)
            faiss.normalize_L2(embeddings)
            index.add(embeddings)

            # Save DIRECTLY to Supabase (no buffer storage)
            await self._save_to_supabase_direct(document_id, embeddings, chunks, file_name, source_info)

            logger.info(f"Created and saved {len(chunks)} vectors directly to Supabase")
            return len(chunks)

        except Exception as e:
            logger.error(f"Error creating embeddings: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to create embeddings: {str(e)}"
            )

    async def search_similar_chunks(self, document_id: str, query: str, top_k: int = TOP_K_RETRIEVAL) -> List[
        Dict[str, Any]]:
        """Search for similar chunks by loading from Supabase on-demand"""
        try:
            logger.info(f"🔍 Searching document {document_id} for: '{query[:50]}...'")

            # Load embeddings and chunks from Supabase
            embeddings, chunks = await self._load_from_supabase_direct(document_id)

            # Create temporary FAISS index for search
            index = faiss.IndexFlatIP(self.dimension)
            faiss.normalize_L2(embeddings)
            index.add(embeddings)

            logger.info(f"📊 Loaded {index.ntotal} vectors from Supabase")

            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])
            query_embedding = query_embedding.astype('float32')
            faiss.normalize_L2(query_embedding)

            # Search similar vectors
            scores, indices = index.search(query_embedding, min(top_k, index.ntotal))

            # Retrieve corresponding chunks
            retrieved_chunks = []
            for i, (score, idx) in enumerate(zip(scores[0], indices[0])):
                if idx != -1:
                    chunk = chunks[idx].copy()
                    chunk["similarity_score"] = float(score)
                    chunk["rank"] = i + 1
                    retrieved_chunks.append(chunk)

            logger.info(f"✅ Retrieved {len(retrieved_chunks)} relevant chunks")
            return retrieved_chunks

        except Exception as e:
            logger.error(f"❌ Error searching chunks: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to search chunks: {str(e)}"
            )

    async def _save_to_supabase_direct(self, document_id: str, embeddings: np.ndarray, chunks: List[Dict[str, Any]],
                                       file_name: str = None, source_info: Dict[str, Any] = None):
        """Save embeddings and chunks directly to Supabase as separate files"""
        try:
            logger.info(f"Saving vector data directly to Supabase for {document_id}")

            # Save embeddings as numpy array
            embeddings_bytes = io.BytesIO()
            np.save(embeddings_bytes, embeddings)
            embeddings_data = embeddings_bytes.getvalue()

            embeddings_path = f"vectors/{document_id}/embeddings.npy"
            await upload_file_to_supabase(embeddings_path, embeddings_data)

            # Save chunks as JSON
            chunks_json = json.dumps(chunks, indent=2)
            chunks_bytes = chunks_json.encode('utf-8')
            chunks_path = f"vectors/{document_id}/chunks.json"
            await upload_file_to_supabase(chunks_path, chunks_bytes)

            # Save enhanced metadata with both filename variants
            metadata = {
                "document_id": document_id,
                "chunks_count": len(chunks),
                "embedding_model": EMBEDDING_MODEL_NAME,
                "dimension": self.dimension,
                "created_at": time.time(),
                "total_characters": sum(chunk["char_count"] for chunk in chunks),
                "file_name": file_name,  # This will be the original readable name
                "original_filename": source_info.get("original_filename", file_name) if source_info else file_name,
                # NEW
                "sanitized_filename": source_info.get("sanitized_filename", file_name) if source_info else file_name,
                # NEW
                "source_info": source_info or {}
            }
            metadata_json = json.dumps(metadata, indent=2)
            metadata_bytes = metadata_json.encode('utf-8')
            metadata_path = f"vectors/{document_id}/metadata.json"
            await upload_file_to_supabase(metadata_path, metadata_bytes)

            logger.info(f"Saved vector data directly to Supabase: {embeddings_path}, {chunks_path}, {metadata_path}")

        except Exception as e:
            logger.error(f"Error saving to Supabase: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to save vector data: {str(e)}"
            )

    async def _load_from_supabase_direct(self, document_id: str) -> Tuple[np.ndarray, List[Dict[str, Any]]]:
        """Load embeddings and chunks directly from Supabase"""
        try:
            logger.info(f"📥 Loading vector data from Supabase for {document_id}")

            # Load embeddings
            embeddings_path = f"vectors/{document_id}/embeddings.npy"
            try:
                embeddings_bytes = await download_file_from_supabase(embeddings_path)
                embeddings_io = io.BytesIO(embeddings_bytes)
                embeddings = np.load(embeddings_io)
            except Exception as e:
                logger.error(f"Failed to load embeddings from {embeddings_path}: {e}")
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=f"Embeddings not found for document {document_id}"
                )

            # Load chunks
            chunks_path = f"vectors/{document_id}/chunks.json"
            try:
                chunks_bytes = await download_file_from_supabase(chunks_path)
                chunks_json = chunks_bytes.decode('utf-8')
                chunks = json.loads(chunks_json)
            except Exception as e:
                logger.error(f"Failed to load chunks from {chunks_path}: {e}")
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=f"Chunks not found for document {document_id}"
                )

            logger.info(f"✅ Loaded {len(embeddings)} embeddings and {len(chunks)} chunks from Supabase")
            return embeddings, chunks

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error loading from Supabase: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to load vector data: {str(e)}"
            )

    async def delete_document_vectors(self, document_id: str) -> bool:
        """Delete document vectors from Supabase"""
        try:
            logger.info(f"🗑️ Deleting vector data for document {document_id}")

            # Delete all vector files
            embeddings_path = f"vectors/{document_id}/embeddings.npy"
            chunks_path = f"vectors/{document_id}/chunks.json"
            metadata_path = f"vectors/{document_id}/metadata.json"

            results = []
            for path in [embeddings_path, chunks_path, metadata_path]:
                try:
                    result = await delete_file_from_supabase(path)
                    results.append(result)
                except Exception as e:
                    logger.warning(f"Failed to delete {path}: {e}")
                    results.append(False)

            success = any(results)
            logger.info(f"✅ Deleted vector data for document {document_id}")
            return success

        except Exception as e:
            logger.error(f"Error deleting document vectors: {e}")
            return False

    async def list_stored_documents(self) -> List[Dict[str, Any]]:
        """List all documents with vector data in Supabase"""
        try:
            # List files specifically in the vectors/ directory
            files = await list_supabase_files(prefix="vectors/")

            documents = []
            seen_doc_ids = set()

            for file_info in files:
                file_path = file_info.get('name', '')

                # Look for metadata.json files in vectors/doc_id/metadata.json pattern
                if file_path.startswith('vectors/') and file_path.endswith('/metadata.json'):
                    try:
                        # Extract document_id from path: vectors/{doc_id}/metadata.json
                        path_parts = file_path.split('/')
                        if len(path_parts) >= 3:
                            document_id = path_parts[1]

                            # Skip if we've already processed this document
                            if document_id in seen_doc_ids:
                                continue
                            seen_doc_ids.add(document_id)

                            # Load metadata
                            metadata_bytes = await download_file_from_supabase(file_path)
                            metadata = json.loads(metadata_bytes.decode('utf-8'))

                            documents.append({
                                "document_id": document_id,
                                "file_name": metadata.get("file_name", "unknown"),
                                "chunks_count": metadata.get("chunks_count", 0),
                                "total_characters": metadata.get("total_characters", 0),
                                "embedding_model": metadata.get("embedding_model", "unknown"),
                                "created_at": metadata.get("created_at", 0),
                                "status": "stored_in_supabase",
                                "supabase_path": f"vectors/{document_id}/"
                            })

                    except Exception as e:
                        logger.warning(f"Failed to process metadata file {file_path}: {e}")
                        continue

            logger.info(f"Found {len(documents)} documents in Supabase vectors/ directory")
            return documents

        except Exception as e:
            logger.error(f"Error listing documents: {e}")
            return []

    async def get_document_chunks(self, document_id: str) -> List[Dict[str, Any]]:
        """Get chunks for a specific document from Supabase"""
        try:
            chunks_path = f"vectors/{document_id}/chunks.json"
            chunks_bytes = await download_file_from_supabase(chunks_path)
            chunks_json = chunks_bytes.decode('utf-8')
            chunks = json.loads(chunks_json)
            return chunks
        except Exception as e:
            logger.error(f"Error getting chunks: {e}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Chunks not found for document {document_id}"
            )

    async def search_across_documents(self, query: str, top_k: int = 10, max_docs: int = 5) -> List[Dict[str, Any]]:
        """Search across all documents stored in Supabase"""
        try:
            logger.info(f"🔍 Global search: '{query[:50]}...' across max {max_docs} documents")

            # Get all available documents
            all_documents = await self.list_stored_documents()

            if not all_documents:
                logger.warning("No documents found in Supabase")
                return []

            # Shortlist documents (simple heuristic: by recent creation + character count)
            shortlisted_docs = sorted(
                all_documents,
                key=lambda x: (x.get('created_at', 0), x.get('total_characters', 0)),
                reverse=True
            )[:max_docs]

            logger.info(
                f"📋 Shortlisted {len(shortlisted_docs)} documents: {[d['document_id'][:8] for d in shortlisted_docs]}")

            # Search each document concurrently
            search_tasks = []
            for doc in shortlisted_docs:
                task = self._search_single_document_with_metadata(doc, query, top_k)
                search_tasks.append(task)

            # Execute searches concurrently
            results = await asyncio.gather(*search_tasks, return_exceptions=True)

            # Merge and sort results
            all_chunks = []
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    logger.warning(f"Search failed for document {shortlisted_docs[i]['document_id']}: {result}")
                    continue
                all_chunks.extend(result)

            # Sort by similarity score and return top_k
            all_chunks.sort(key=lambda x: x['similarity_score'], reverse=True)
            final_results = all_chunks[:top_k]

            logger.info(
                f"✅ Global search completed: {len(final_results)} results from {len(shortlisted_docs)} documents")
            return final_results

        except Exception as e:
            logger.error(f"Error in global search: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to perform global search: {str(e)}"
            )

    async def search_filtered_documents(self, query: str, document_ids: List[str], top_k: int = 10) -> List[
        Dict[str, Any]]:
        """Search across specific documents by ID or filename"""
        try:
            logger.info(f"🔍 Filtered search: '{query[:50]}...' across {len(document_ids)} specified documents")

            # Get all available documents
            all_documents = await self.list_stored_documents()

            if not all_documents:
                logger.warning("No documents found in Supabase")
                return []

            # Filter documents by ID or filename
            filtered_docs = []
            for doc in all_documents:
                doc_id = doc["document_id"]
                file_name = doc.get("file_name", "unknown")

                # Check if document matches any of the provided IDs or filenames
                if (doc_id in document_ids or
                        file_name in document_ids or
                        any(doc_id.startswith(did) for did in document_ids if len(did) >= 8) or  # Partial ID match
                        any(file_name.lower() == did.lower() for did in
                            document_ids)):  # Case-insensitive filename match
                    filtered_docs.append(doc)

            if not filtered_docs:
                logger.warning(f"No documents found matching IDs/filenames: {document_ids}")
                return []

            logger.info(f"📋 Found {len(filtered_docs)} matching documents: {[d['file_name'] for d in filtered_docs]}")

            # Search each filtered document concurrently
            search_tasks = []
            for doc in filtered_docs:
                task = self._search_single_document_with_metadata(doc, query, top_k)
                search_tasks.append(task)

            # Execute searches concurrently
            results = await asyncio.gather(*search_tasks, return_exceptions=True)

            # Merge and sort results
            all_chunks = []
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    logger.warning(f"Search failed for document {filtered_docs[i]['document_id']}: {result}")
                    continue
                all_chunks.extend(result)

            # Sort by similarity score and return top_k
            all_chunks.sort(key=lambda x: x['similarity_score'], reverse=True)
            final_results = all_chunks[:top_k]

            logger.info(
                f"✅ Filtered search completed: {len(final_results)} results from {len(filtered_docs)} documents")
            return final_results

        except Exception as e:
            logger.error(f"Error in filtered document search: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to perform filtered document search: {str(e)}"
            )

    async def _search_single_document_with_metadata(self, doc_metadata: Dict[str, Any], query: str, top_k: int) -> List[
        Dict[str, Any]]:
        """Search a single document and add document metadata to results"""
        try:
            document_id = doc_metadata['document_id']

            # Load embeddings and chunks
            embeddings, chunks = await self._load_from_supabase_direct(document_id)

            # Create temporary FAISS index
            index = faiss.IndexFlatIP(self.dimension)
            faiss.normalize_L2(embeddings)
            index.add(embeddings)

            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])
            query_embedding = query_embedding.astype('float32')
            faiss.normalize_L2(query_embedding)

            # Search
            scores, indices = index.search(query_embedding, min(top_k, index.ntotal))

            # Build results with document metadata
            results = []
            for score, idx in zip(scores[0], indices[0]):
                if idx != -1 and score > 0.3:  # Minimum similarity threshold
                    chunk = chunks[idx].copy()
                    chunk.update({
                        "similarity_score": float(score),
                        "document_id": document_id,
                        "file_name": doc_metadata.get('file_name', 'unknown'),
                        "chunk_preview": chunk["text"][:150] + "..." if len(chunk["text"]) > 150 else chunk["text"]
                    })
                    results.append(chunk)

            logger.info(f"📄 Found {len(results)} relevant chunks in {doc_metadata.get('file_name', document_id[:8])}")
            return results

        except Exception as e:
            logger.warning(f"Failed to search document {document_id}: {e}")
            return []


class RAGLLMService:
    """Enhanced LLM service with RAG capabilities"""

    def __init__(self):
        self.model = None
        self.model_name = None
        self.safety_settings = {
            HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
        }
        self._initialize_model()

    def _initialize_model(self):
        """Initialize Gemini model with optimized settings"""
        try:
            model_variants = [
                'gemini-1.5-flash-latest',
                'gemini-1.5-flash',
                'gemini-1.0-pro-latest',
                'gemini-1.0-pro'
            ]

            for model_name in model_variants:
                try:
                    test_model = genai.GenerativeModel(model_name)
                    test_response = test_model.generate_content(
                        "Test message. Respond with 'OK'.",
                        generation_config=genai.types.GenerationConfig(
                            temperature=0.1,
                            max_output_tokens=10,
                        ),
                        safety_settings=self.safety_settings
                    )

                    if test_response and test_response.text:
                        self.model = test_model
                        self.model_name = model_name
                        logger.info(f"✅ Successfully initialized {model_name}")
                        return

                except Exception as e:
                    logger.warning(f"Failed to initialize {model_name}: {e}")
                    continue

            raise Exception("No working Gemini model found")

        except Exception as e:
            logger.error(f"❌ Failed to initialize Gemini model: {e}")
            self.model = None
            self.model_name = None

    async def generate_rag_answer(self, question: str, relevant_chunks: List[Dict[str, Any]], document_id: str) -> Dict[
        str, Any]:
        """Generate answer using RAG approach - FIXED VERSION"""
        if not self.model:
            return {
                "answer": "LLM service not available",
                "confidence": 0.0,
                "sources": [],
                "chunks_retrieved": 0
            }

        try:
            # Log retrieval info
            logger.info(f"🔍 Retrieved {len(relevant_chunks)} chunks for question: {question[:50]}...")

            if not relevant_chunks:
                logger.warning(f"⚠️ No context retrieved for document: {document_id}")
                return {
                    "answer": "No relevant context found in the document to answer this question.",
                    "confidence": 0.0,
                    "sources": [],
                    "chunks_retrieved": 0
                }

            # Log top retrieved chunks for debugging
            for i, chunk in enumerate(relevant_chunks[:3]):
                logger.info(
                    f"📄 Chunk {i + 1} (score: {chunk.get('similarity_score', 0):.3f}): {chunk['text'][:100]}...")

            # Construct context from relevant chunks
            context = self._build_context(relevant_chunks)

            # Create RAG prompt
            prompt = self._create_rag_prompt(question, context)

            # Generate response
            response = self.model.generate_content(
                prompt,
                safety_settings=self.safety_settings,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.2,
                    max_output_tokens=500,
                    top_p=0.8,
                    top_k=40
                )
            )

            if response and response.text:
                answer_text = response.text.strip()
                confidence = self._estimate_confidence(relevant_chunks)

                # Extract sources
                sources = [
                    {
                        "chunk_id": chunk["chunk_id"],
                        "similarity_score": chunk["similarity_score"],
                        "preview": chunk["text"][:200] + "..." if len(chunk["text"]) > 200 else chunk["text"]
                    }
                    for chunk in relevant_chunks[:3]  # Top 3 sources
                ]

                logger.info(f"✅ Generated answer with confidence {confidence:.2f}")

                return {
                    "answer": answer_text,
                    "confidence": confidence,
                    "sources": sources,
                    "chunks_retrieved": len(relevant_chunks)
                }
            else:
                logger.error("❌ Gemini returned empty response")
                return {
                    "answer": "Unable to generate response - empty response from LLM",
                    "confidence": 0.0,
                    "sources": [],
                    "chunks_retrieved": len(relevant_chunks)
                }

        except Exception as e:
            logger.error(f"❌ Error generating RAG answer: {e}")
            return {
                "answer": f"Error generating answer: {str(e)}",
                "confidence": 0.0,
                "sources": [],
                "chunks_retrieved": len(relevant_chunks)
            }

    def _build_context(self, chunks: List[Dict[str, Any]]) -> str:
        """Build context from retrieved chunks"""
        context_parts = []
        total_length = 0

        for i, chunk in enumerate(chunks):
            chunk_text = f"[Context {i + 1}]\n{chunk['text']}\n"

            if total_length + len(chunk_text) > MAX_CONTEXT_LENGTH:
                break

            context_parts.append(chunk_text)
            total_length += len(chunk_text)

        return "\n".join(context_parts)

    def _create_rag_prompt(self, question: str, context: str) -> str:
        return f"""You are an expert assistant trained to answer questions based strictly on the following insurance policy content.

    DOCUMENT EXTRACT:
    {context}

    QUESTION:
    {question}

    RESPONSE INSTRUCTIONS:
    - Answer using only the document content.
    - Provide complete, factual, and specific responses.
    - Include exact terms, durations, limits, and conditions as stated in the document.
    - If the document lists multiple relevant points or sub-clauses, include all of them.
    - Format clearly in professional language, as if responding to a client or regulator.
    - Do NOT add disclaimers, assumptions, or refer to "the context".
    - If the information is not present, respond with: "The provided document does not contain sufficient information to answer this question."

    FINAL ANSWER:"""

    def _estimate_confidence(self, chunks: List[Dict[str, Any]]) -> float:
        """Estimate confidence based on similarity scores"""
        if not chunks:
            return 0.0

        # Use weighted average of top chunks
        weights = [1.0, 0.8, 0.6, 0.4, 0.2]
        total_score = 0.0
        total_weight = 0.0

        for i, chunk in enumerate(chunks[:5]):
            weight = weights[i] if i < len(weights) else 0.1
            total_score += chunk["similarity_score"] * weight
            total_weight += weight

        confidence = total_score / total_weight if total_weight > 0 else 0.0
        return min(confidence, 1.0)

class DocumentIngestRequest(BaseModel):
    source: str  # URL, Google Drive link, or direct document link
    source_type: Optional[str] = "auto"  # auto, url, google_drive, direct_link
    folder_name: Optional[str] = None  # For organizing multiple documents

    @field_validator('source')
    @classmethod
    def validate_source(cls, v):
        if not v or not v.strip():
            raise ValueError("Source cannot be empty")
        return v.strip()


# Initialize services
vector_store = VectorStore()
rag_llm_service = RAGLLMService()


def is_valid_document_url(url: str) -> bool:
    """Check if URL likely points to a valid document format"""
    url_lower = url.lower()

    # Valid extensions
    valid_extensions = ['.pdf', '.docx', '.doc', '.txt', '.eml']

    # Check if URL contains valid extension
    if any(ext in url_lower for ext in valid_extensions):
        return True

    # For Google Drive URLs, assume valid (will be validated during download)
    if 'drive.google.com' in url_lower or 'docs.google.com' in url_lower:
        return True

    # For other URLs without clear extensions, we'll validate during download
    return True


async def _ingest_single_document(request: DocumentIngestRequest) -> Dict[str, Any]:
    """Ingest a single document from URL with enhanced validation"""
    try:
        # Validate file type before processing (for URLs)
        if request.source.startswith(('http://', 'https://')):
            if not is_valid_document_url(request.source):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="URL does not appear to point to a valid document format. Supported: PDF, DOCX, DOC, TXT, EML"
                )

        # Extract text and get document info
        text, document_id, original_filename, sanitized_filename = await DocumentProcessor.process_document_from_source(
            request.source, request.source_type
        )

        # Check if document already exists in Supabase
        try:
            await vector_store._load_from_supabase_direct(document_id)
            logger.info(f"Found existing vector data for document {document_id}")
            return {
                "document_id": document_id,
                "filename": original_filename,
                "original_filename": original_filename,
                "sanitized_filename": sanitized_filename,
                "status": "already_processed",
                "source": request.source,
                "message": "Document already exists in vector store"
            }
        except HTTPException:
            # Document not found, create new one
            logger.info(f"Creating new vector index for document {document_id}")

            # Create chunks
            chunks = DocumentProcessor.intelligent_chunking(text)

            # Prepare source info with both filenames
            source_info = {
                "source_url": request.source,
                "source_type": request.source_type or "auto",
                "folder_name": request.folder_name,
                "original_filename": original_filename,
                "sanitized_filename": sanitized_filename,
                "ingested_at": time.time()
            }

            # Create embeddings and store
            chunks_created = await vector_store.create_embeddings(
                document_id, chunks, original_filename, source_info
            )

            return {
                "document_id": document_id,
                "filename": original_filename,
                "original_filename": original_filename,
                "sanitized_filename": sanitized_filename,
                "status": "processed",
                "chunks_created": chunks_created,
                "source": request.source,
                "source_info": source_info,
                "message": f"Document processed successfully with {chunks_created} chunks"
            }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing single document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process document: {str(e)}"
        )

# Replace the root endpoint (around line 280) with this enhanced version:

@app.get("/")
async def list_all_documents():
    """Root endpoint - List all available documents"""
    try:
        logger.info("Fetching all documents from Supabase...")

        # Get all documents from vector store
        documents = await vector_store.list_stored_documents()

        # Format documents in the requested format
        formatted_docs = []
        for doc in documents:
            formatted_docs.append({
                "document_id": doc["document_id"],
                "file_name": doc.get("file_name", "unknown"),
                "chunks_count": doc.get("chunks_count", 0),
                "created_at": int(doc.get("created_at", 0))
            })

        return formatted_docs

    except Exception as e:
        logger.error(f"Error fetching documents: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to fetch documents: {str(e)}"
        )


# API Routes
@app.get("/api/v1/health")
async def health_check():
    """Enhanced health check with Supabase status for Cloud Run"""
    supabase_status = "unknown"
    try:
        supabase_manager = get_supabase_manager()
        # Quick test Supabase connection
        await supabase_manager.list_files()
        supabase_status = "connected"
    except Exception as e:
        supabase_status = f"error: {str(e)[:100]}"  # Truncate long error messages

    return {
        "status": "healthy",
        "timestamp": time.time(),
        "environment": os.getenv("ENVIRONMENT", "development"),
        "platform": "google-cloud-run",
        "services": {
            "gemini_available": rag_llm_service.model is not None,
            "gemini_model": rag_llm_service.model_name,
            "embedding_model": EMBEDDING_MODEL_NAME,
            "vector_store_ready": True,
            "supabase_status": supabase_status
        },
        "configuration": {
            "chunk_size": CHUNK_SIZE,
            "chunk_overlap": CHUNK_OVERLAP,
            "top_k_retrieval": TOP_K_RETRIEVAL,
            "max_context_length": MAX_CONTEXT_LENGTH,
            "supabase_bucket": os.getenv("SUPABASE_BUCKET", "documents"),
            "storage_mode": "FULL_SUPABASE_ONLY"
        },
        "memory_usage": {
            "cached_documents": 0,  # No more caching
            "cached_chunks": 0,     # No more buffering
            "storage_mode": "direct_supabase_only"
        }
    }


@app.post("/api/v1/documents/upload")
async def upload_document(
        file: UploadFile = File(...),
        token: str = Depends(verify_token)
):
    """Upload and process document for RAG using Supabase Storage - FIXED VERSION"""
    logger.info("=" * 100)
    logger.info("📤 DOCUMENT UPLOAD DEBUG SESSION STARTED")
    logger.info("=" * 100)

    try:
        # Log upload details
        logger.info(f"📁 Original filename: {file.filename}")
        logger.info(f"📋 Content type: {file.content_type}")
        logger.info(f"🆔 Generating document ID...")

        # Generate document ID
        document_id = str(uuid.uuid4())
        logger.info(f"✅ Generated document ID: {document_id}")

        # Create Supabase file path with sanitized filename but preserve original
        timestamp = int(time.time())
        sanitized_filename = DocumentProcessor.sanitize_filename(file.filename)
        supabase_file_path = f"documents/{document_id}_{timestamp}_{sanitized_filename}"
        logger.info(f"🗂️ Supabase file path: {supabase_file_path}")

        # Read file content
        logger.info("📖 Reading file content...")
        content = await file.read()
        logger.info(f"✅ File content read: {len(content)} bytes")

        # CRITICAL: Upload original file to Supabase FIRST
        logger.info("🚀 UPLOADING FILE TO SUPABASE...")
        logger.info(f"   📁 File path: {supabase_file_path}")
        logger.info(f"   📊 File size: {len(content)} bytes")
        logger.info(f"   🪣 Target bucket: {os.getenv('SUPABASE_BUCKET', 'documents')}")

        try:
            uploaded_path = await upload_file_to_supabase(supabase_file_path, content)
            logger.info(f"✅ SUPABASE UPLOAD SUCCESS: {uploaded_path}")
        except Exception as upload_error:
            logger.error(f"❌ SUPABASE UPLOAD FAILED: {upload_error}")
            logger.error(f"❌ Upload error details: {type(upload_error).__name__}: {str(upload_error)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to upload file to Supabase: {str(upload_error)}"
            )

        # Get public URL (now async)
        logger.info("🔗 Getting public URL...")
        try:
            public_url = await get_public_url(uploaded_path)
            logger.info(f"✅ Public URL generated: {public_url}")
        except Exception as url_error:
            logger.warning(f"⚠️ Could not generate public URL: {url_error}")
            public_url = None

        # Process document content
        logger.info("📤 Processing document text...")
        processor = DocumentProcessor()
        text = await processor.process_uploaded_file(content, file.filename)
        logger.info(f"✅ Text extracted: {len(text)} characters")

        # Create intelligent chunks
        logger.info("🧩 Creating intelligent chunks...")
        chunks = processor.intelligent_chunking(text)
        logger.info(f"✅ Created {len(chunks)} chunks")

        # Create embeddings and vector index (stored in Supabase) - UPDATED
        logger.info("🧠 Creating embeddings and saving to Supabase...")

        # Create source info with enhanced metadata - UPDATED
        source_info = {
            "source_type": "file_upload",
            "upload_timestamp": timestamp,
            "original_content_type": file.content_type,
            "original_filename": file.filename,  # Keep original readable name
            "sanitized_filename": sanitized_filename,  # Safe storage name
            "supabase_path": uploaded_path,
            "file_size_bytes": len(content),
            "ingested_at": time.time(),
            "upload_method": "multipart_form"
        }

        chunks_created = await vector_store.create_embeddings(document_id, chunks, file.filename, source_info)
        logger.info(f"✅ Created and stored {chunks_created} embeddings in Supabase")

        logger.info("=" * 100)
        logger.info("✅ DOCUMENT UPLOAD COMPLETED SUCCESSFULLY!")
        logger.info(f"   📄 Document ID: {document_id}")
        logger.info(f"   📁 Filename: {file.filename}")
        logger.info(f"   📁 Supabase path: {uploaded_path}")
        logger.info(f"   🧩 Chunks created: {chunks_created}")
        logger.info(f"   🔗 Public URL: {public_url or 'Not generated'}")
        logger.info("=" * 100)

        return DocumentUploadResponse(
            document_id=document_id,
            filename=file.filename,
            status="processed",
            chunks_created=chunks_created,
            message=f"Document processed successfully with {chunks_created} chunks (stored in Supabase)",
            supabase_path=uploaded_path,
            public_url=public_url if public_url else None
        )

    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        logger.error("=" * 100)
        logger.error("❌ DOCUMENT UPLOAD FAILED!")
        logger.error(f"❌ Error: {type(e).__name__}: {str(e)}")
        logger.error(f"❌ Full traceback:")
        import traceback
        logger.error(traceback.format_exc())
        logger.error("=" * 100)

        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process document: {str(e)}"
        )


@app.post("/api/v1/documents/ingest")
async def ingest_documents_unified(
        request: DocumentIngestRequest,
        token: str = Depends(verify_token)
):
    """Unified endpoint to ingest single documents, folders, or nested folder structures"""
    try:
        logger.info(f"📥 Unified ingestion from source: {request.source}")

        # Check if it's a Google Drive folder
        if GoogleDriveProcessor.is_google_drive_link(request.source) and '/folders/' in request.source:
            return await _ingest_google_drive_folder_recursive(request)
        else:
            # For non-folder sources, process as single document
            logger.info("Source is not a folder, processing as single document")
            return await _ingest_single_document(request)

    except Exception as e:
        logger.error(f"Error in unified ingestion: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to ingest documents: {str(e)}"
        )


async def _ingest_google_drive_folder_recursive(request: DocumentIngestRequest) -> Dict[str, Any]:
    """Enhanced Google Drive folder processing with recursive folder support and file filtering"""
    try:
        # Get list of files in folder
        files = await GoogleDriveProcessor.get_drive_folder_files_recursive(request.source)

        if not files:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="No files found in Google Drive folder or folder is not publicly accessible"
            )

        # Filter valid document files
        valid_files = _filter_valid_document_files(files)

        if not valid_files:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No valid document files found in folder. Supported formats: PDF, DOCX, DOC, TXT, EML"
            )

        logger.info(f"📁 Processing {len(valid_files)} valid documents from {len(files)} total files")

        results = []
        successful = 0
        failed = 0
        skipped = 0

        for i, file_info in enumerate(valid_files):
            file_name = file_info.get("original_filename", file_info.get("filename", "unknown"))
            folder_path = file_info.get("folder_path", "")
            logger.info(f"Processing file {i + 1}/{len(valid_files)}: {folder_path}/{file_name}")

            try:
                # Create individual ingest request with folder context
                file_request = DocumentIngestRequest(
                    source=file_info["download_url"],
                    source_type="google_drive",
                    folder_name=f"{request.folder_name or 'Google Drive'}/{folder_path}" if folder_path else request.folder_name or "Google Drive"
                )

                # Ingest individual file with timeout
                result = await asyncio.wait_for(
                    _ingest_single_document(file_request),
                    timeout=120.0  # 2 minute timeout per file
                )

                # Override filename with the real Drive filename and add folder info
                result["filename"] = file_name
                result["folder_path"] = folder_path
                result["drive_metadata"] = {
                    "drive_file_id": file_info.get("file_id"),
                    "original_filename": file_name,
                    "folder_path": folder_path,
                    "file_type": file_info.get("file_type", "unknown")
                }
                results.append(result)

                if result["status"] in ["processed", "already_processed"]:
                    successful += 1
                else:
                    failed += 1

                logger.info(f"✅ Successfully processed: {folder_path}/{file_name}")

            except asyncio.TimeoutError:
                logger.warning(f"⏱️ Timeout processing file: {folder_path}/{file_name}")
                results.append({
                    "filename": file_name,
                    "folder_path": folder_path,
                    "status": "timeout",
                    "error": "File processing timed out",
                    "original_file_info": file_info
                })
                failed += 1

            except HTTPException as he:
                # Handle specific HTTP errors more gracefully
                if he.status_code == 400 and "HTML instead of document" in str(he.detail):
                    logger.warning(f"⭐ Skipping inaccessible file: {folder_path}/{file_name}")
                    results.append({
                        "filename": file_name,
                        "folder_path": folder_path,
                        "status": "skipped",
                        "error": "File not publicly accessible",
                        "original_file_info": file_info
                    })
                    skipped += 1
                elif he.status_code == 422 and "EOF marker not found" in str(he.detail):
                    logger.warning(f"⭐ Skipping corrupted file: {folder_path}/{file_name}")
                    results.append({
                        "filename": file_name,
                        "folder_path": folder_path,
                        "status": "skipped",
                        "error": "File appears to be corrupted or incomplete",
                        "original_file_info": file_info
                    })
                    skipped += 1
                else:
                    logger.error(f"❌ HTTP error processing {folder_path}/{file_name}: {he.detail}")
                    results.append({
                        "filename": file_name,
                        "folder_path": folder_path,
                        "status": "failed",
                        "error": str(he.detail),
                        "original_file_info": file_info
                    })
                    failed += 1

            except Exception as e:
                logger.error(f"❌ Failed to process file {folder_path}/{file_name}: {e}")
                results.append({
                    "filename": file_name,
                    "folder_path": folder_path,
                    "status": "failed",
                    "error": str(e),
                    "original_file_info": file_info
                })
                failed += 1

            # Rate limiting between files
            if i < len(valid_files) - 1:  # Don't sleep after last file
                await asyncio.sleep(2)

        return {
            "folder_url": request.source,
            "total_files_found": len(files),
            "valid_document_files": len(valid_files),
            "successful": successful,
            "failed": failed,
            "skipped": skipped,
            "folder_name": request.folder_name or "Google Drive Folder",
            "results": results,
            "message": f"Processed {successful}/{len(valid_files)} valid documents successfully, {skipped} skipped, {failed} failed",
            "note": "Only valid document formats (PDF, DOCX, DOC, TXT, EML) are processed. Folders are processed recursively.",
            "processing_info": {
                "recursive_folders": True,
                "file_filtering": True,
                "supported_formats": ["PDF", "DOCX", "DOC", "TXT", "EML"]
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing Google Drive folder: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process Google Drive folder: {str(e)}"
        )


def _filter_valid_document_files(files: List[Dict[str, str]]) -> List[Dict[str, str]]:
    """Filter files to only include valid document formats"""
    valid_extensions = {'.pdf', '.docx', '.doc', '.txt', '.eml'}
    valid_files = []

    for file_info in files:
        filename = file_info.get("original_filename", file_info.get("filename", "")).lower()

        # Check if file has a valid document extension
        is_valid = any(filename.endswith(ext) for ext in valid_extensions)

        if is_valid:
            # Add file type info for better tracking
            file_type = "unknown"
            for ext in valid_extensions:
                if filename.endswith(ext):
                    file_type = ext[1:]  # Remove the dot
                    break

            file_info["file_type"] = file_type
            valid_files.append(file_info)
        else:
            logger.info(f"🚫 Skipping invalid file type: {filename}")

    logger.info(f"📋 Filtered to {len(valid_files)} valid documents from {len(files)} total files")
    return valid_files


@app.post("/api/v1/query/global")
async def global_query(
        request: GlobalQueryRequest,
        token: str = Depends(verify_token)
):
    """Global search across all documents or filtered documents with RAG answer generation"""
    try:
        if request.document_ids:
            logger.info(f"🌐 Filtered query: '{request.query[:50]}...' across {len(request.document_ids)} specified documents")
            # Search only specified documents
            relevant_chunks = await vector_store.search_filtered_documents(
                request.query,
                request.document_ids,
                top_k=request.top_k
            )
            search_type = "filtered"
            documents_searched_info = f"Searched {len(request.document_ids)} specified documents"
        else:
            logger.info(f"🌐 Global query: '{request.query[:50]}...' across all documents")
            # Search across all documents (existing behavior)
            relevant_chunks = await vector_store.search_across_documents(
                request.query,
                top_k=request.top_k,
                max_docs=request.max_docs
            )
            search_type = "global"
            documents_searched_info = f"Searched up to {request.max_docs} documents"

        if not relevant_chunks:
            search_scope = f"specified documents ({request.document_ids})" if request.document_ids else "all documents"
            return {
                "answer": f"No relevant information found in {search_scope}.",
                "sources": [],
                "search_type": search_type,
                "documents_searched": 0
            }

        # Generate RAG answer using top chunks from search
        answer_data = await rag_llm_service.generate_rag_answer(
            request.query,
            relevant_chunks,
            "filtered_search" if request.document_ids else "global_search"
        )

        # Format sources with document info
        sources = []
        for chunk in relevant_chunks[:5]:  # Top 5 sources
            sources.append({
                "document_id": chunk["document_id"],
                "file_name": chunk.get("file_name", "unknown"),
                "chunk_preview": chunk.get("chunk_preview", chunk["text"][:150] + "..."),
                "similarity_score": chunk["similarity_score"]
            })

        unique_doc_count = len(set(chunk["document_id"] for chunk in relevant_chunks))

        return {
            "answer": answer_data.get("answer", "No answer generated"),
            "sources": sources,
            "query": request.query,
            "search_type": search_type,
            "chunks_searched": len(relevant_chunks),
            "documents_searched": unique_doc_count,
            "filter_applied": request.document_ids is not None,
            "filtered_documents": request.document_ids if request.document_ids else None,
            "search_info": documents_searched_info
        }

    except Exception as e:
        logger.error(f"Error in global query: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process global query: {str(e)}"
        )


@app.post(
    "/api/v1/hackrx/run",
    responses={
        400: {"model": ErrorResponse},
        401: {"model": ErrorResponse},
        422: {"model": ErrorResponse},
        503: {"model": ErrorResponse}
    }
)
async def process_document_qa_rag(
        request: DocumentQARequest,
        token: str = Depends(verify_token)
):
    """
    RAG-powered document QA endpoint - Returns detailed answers with sources and metadata
    """
    try:
        document_id = None

        # Process document if source provided
        if request.documents and not request.document_id:
            logger.info(f"Processing document from source: {request.documents}")

            # Extract text and get document ID with both filenames - UPDATED
            text, document_id, original_filename, sanitized_filename = await DocumentProcessor.process_document_from_source(
                request.documents)

            # Check if we already have this document processed in Supabase
            try:
                # Try to load from Supabase to see if it exists
                await vector_store._load_from_supabase_direct(document_id)
                logger.info(f"Found existing vector data in Supabase for document {document_id}")
            except HTTPException:
                # Document not found in Supabase, create new one
                logger.info(f"Creating new vector index for document {document_id}")
                chunks = DocumentProcessor.intelligent_chunking(text)

                # Create source info with both filenames - UPDATED
                source_info = {
                    "source_url": request.documents,
                    "source_type": "api_request",
                    "processed_via": "hackrx_endpoint",
                    "original_filename": original_filename,
                    "sanitized_filename": sanitized_filename,
                    "ingested_at": time.time()
                }

                await vector_store.create_embeddings(document_id, chunks, original_filename, source_info)
                logger.info(f"Created new vector index in Supabase for document {document_id}")

        elif request.document_id:
            document_id = request.document_id
            logger.info(f"Using pre-processed document: {document_id}")

        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Either 'documents' URL/Supabase path or 'document_id' must be provided"
            )

        # Process questions using RAG - Return detailed responses like global query
        detailed_answers = []

        for i, question in enumerate(request.questions):
            logger.info(f"Processing question {i + 1}/{len(request.questions)}: {question[:50]}...")

            # Retrieve relevant chunks from Supabase-stored vectors
            relevant_chunks = await vector_store.search_similar_chunks(
                document_id, question, TOP_K_RETRIEVAL
            )

            # Generate answer using RAG
            answer_data = await rag_llm_service.generate_rag_answer(
                question, relevant_chunks, document_id
            )

            # Format sources with document info (similar to global query)
            sources = []
            for chunk in relevant_chunks[:5]:  # Top 5 sources
                sources.append({
                    "document_id": document_id,
                    "chunk_id": chunk["chunk_id"],
                    "similarity_score": chunk["similarity_score"],
                    "chunk_preview": chunk["text"][:150] + "..." if len(chunk["text"]) > 150 else chunk["text"],
                    "rank": chunk.get("rank", sources.__len__() + 1)
                })

            # Create detailed answer object
            detailed_answer = {
                "question": question,
                "answer": answer_data.get("answer", "No answer generated"),
                "confidence": answer_data.get("confidence", 0.0),
                "sources": sources,
                "chunks_retrieved": len(relevant_chunks),
                "retrieval_info": {
                    "top_similarity_score": relevant_chunks[0]["similarity_score"] if relevant_chunks else 0.0,
                    "avg_similarity_score": sum(chunk["similarity_score"] for chunk in relevant_chunks) / len(relevant_chunks) if relevant_chunks else 0.0,
                    "chunks_above_threshold": len([c for c in relevant_chunks if c["similarity_score"] > 0.5])
                }
            }

            detailed_answers.append(detailed_answer)

            logger.info(f"Generated detailed answer {i + 1} with {len(sources)} sources")

            # Rate limiting for free tier
            await asyncio.sleep(0.5)

        logger.info("Successfully processed all questions using RAG with detailed responses")

        # Return detailed format similar to global query
        return {
            "answers": detailed_answers,
            "document_id": document_id,
            "total_questions": len(request.questions),
            "search_type": "document_specific",
            "processing_info": {
                "document_source": request.documents or "pre_processed",
                "total_chunks_searched": sum(len(ans["sources"]) for ans in detailed_answers),
                "avg_confidence": sum(ans["confidence"] for ans in detailed_answers) / len(detailed_answers) if detailed_answers else 0.0
            },
            "storage_mode": "supabase_direct"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in RAG document QA: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Internal server error: {str(e)}"
        )

@app.get("/api/v1/documents/{document_id}/chunks")
async def get_document_chunks(
        document_id: str,
        token: str = Depends(verify_token)
):
    """Get chunks for a specific document from Supabase - NO BUFFER"""
    try:
        # Load chunks directly from Supabase
        chunks = await vector_store.get_document_chunks(document_id)

        return {
            "document_id": document_id,
            "total_chunks": len(chunks),
            "storage_source": "supabase_direct",
            "chunks": [
                {
                    "chunk_id": chunk["chunk_id"],
                    "preview": chunk["text"][:200] + "..." if len(chunk["text"]) > 200 else chunk["text"],
                    "char_count": chunk["char_count"],
                    "sentence_count": chunk["sentence_count"]
                }
                for chunk in chunks
            ]
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error retrieving chunks: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to retrieve chunks: {str(e)}"
        )

@app.post("/api/v1/documents/{document_id}/search")
async def search_document_chunks(
        document_id: str,
        query: str,
        top_k: int = 5,
        token: str = Depends(verify_token)
):
    """Search for relevant chunks in a specific document stored in Supabase"""
    try:
        relevant_chunks = await vector_store.search_similar_chunks(
            document_id, query, min(top_k, 20)  # Limit to max 20
        )

        return {
            "document_id": document_id,
            "query": query,
            "results_count": len(relevant_chunks),
            "storage_source": "supabase",
            "chunks": [
                {
                    "chunk_id": chunk["chunk_id"],
                    "similarity_score": chunk["similarity_score"],
                    "rank": chunk["rank"],
                    "text": chunk["text"],
                    "char_count": chunk["char_count"]
                }
                for chunk in relevant_chunks
            ]
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error searching chunks: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to search chunks: {str(e)}"
        )


@app.delete("/api/v1/documents/{document_id}")
async def delete_document(
        document_id: str,
        token: str = Depends(verify_token)
):
    """Delete a document and its associated vector data from Supabase"""
    try:
        # Delete from Supabase (this will also clear memory cache)
        success = await vector_store.delete_document_vectors(document_id)

        return {
            "message": f"Document {document_id} deleted from Supabase storage",
            "document_id": document_id,
            "success": success,
            "note": "Original document file may still exist in Supabase. Use /storage endpoint to delete it."
        }

    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete document: {str(e)}"
        )


@app.delete("/api/v1/documents/{document_id}/storage")
async def delete_document_from_storage(
        document_id: str,
        supabase_path: str,
        token: str = Depends(verify_token)
):
    """Delete document from both vector store and original file from Supabase storage"""
    try:
        # Remove vector data
        vector_success = await vector_store.delete_document_vectors(document_id)

        # Remove original file from Supabase storage
        file_success = await delete_file_from_supabase(supabase_path)

        return {
            "message": f"Document {document_id} deleted from Supabase storage",
            "document_id": document_id,
            "vector_data_deleted": vector_success,
            "original_file_deleted": file_success,
            "supabase_path": supabase_path
        }

    except Exception as e:
        logger.error(f"Error deleting document from storage: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete document from storage: {str(e)}"
        )


@app.get("/api/v1/documents")
async def list_documents(token: str = Depends(verify_token)):
    """List all processed documents stored in Supabase"""
    try:
        # Get documents from Supabase
        documents = await vector_store.list_stored_documents()

        return {
            "total_documents": len(documents),
            "storage_mode": "supabase_only",
            "documents": documents
        }

    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to list documents: {str(e)}"
        )


@app.get("/api/v1/storage/files")
async def list_storage_files(token: str = Depends(verify_token)):
    """List all files in Supabase storage with enriched metadata"""
    try:
        # Use the new list function
        files = await list_supabase_files()

        # Separate by file type and enrich with metadata
        document_files = []
        vector_files = []
        other_files = []
        enriched_documents = {}

        # First pass: collect all files and identify vector documents
        for file_info in files:
            file_path = file_info.get('name', '')
            if file_path.startswith('documents/'):
                document_files.append(file_info)
            elif file_path.startswith('vectors/'):
                vector_files.append(file_info)

                # Check if this is a metadata file and extract document_id
                if file_path.endswith('/metadata.json'):
                    path_parts = file_path.split('/')
                    if len(path_parts) >= 3:
                        document_id = path_parts[1]
                        try:
                            # Load and parse metadata
                            metadata_bytes = await download_file_from_supabase(file_path)
                            metadata = json.loads(metadata_bytes.decode('utf-8'))
                            enriched_documents[document_id] = {
                                **file_info,
                                "document_id": document_id,
                                "file_name": metadata.get("file_name", "unknown"),
                                "chunks_count": metadata.get("chunks_count", 0),
                                "total_characters": metadata.get("total_characters", 0),
                                "embedding_model": metadata.get("embedding_model", "unknown"),
                                "created_at": int(metadata.get("created_at", 0)),
                                "original_filename": metadata.get("original_filename"),
                                "sanitized_filename": metadata.get("sanitized_filename"),
                                "source_info": metadata.get("source_info", {}),
                                "storage_path": f"vectors/{document_id}/",
                                "has_metadata": True
                            }
                        except Exception as e:
                            logger.warning(f"Failed to load metadata for {document_id}: {e}")
            else:
                other_files.append(file_info)

        # Second pass: enrich vector files with metadata
        enriched_vector_files = []
        for file_info in vector_files:
            file_path = file_info.get('name', '')
            # Extract document_id from vector file path
            path_parts = file_path.split('/')
            if len(path_parts) >= 3:
                document_id = path_parts[1]
                if document_id in enriched_documents:
                    # Add file type information
                    file_type = "unknown"
                    if file_path.endswith('.npy'):
                        file_type = "embeddings"
                    elif file_path.endswith('.json'):
                        if file_path.endswith('metadata.json'):
                            file_type = "metadata"
                        elif file_path.endswith('chunks.json'):
                            file_type = "chunks"

                    enriched_file = {
                        **file_info,
                        "document_id": document_id,
                        "file_type": file_type,
                        "parent_document": enriched_documents[document_id]
                    }
                    enriched_vector_files.append(enriched_file)
                else:
                    # No metadata found, add basic info
                    enriched_vector_files.append({
                        **file_info,
                        "document_id": document_id,
                        "file_type": "orphaned",
                        "parent_document": None
                    })

        # Create document summary (similar to root endpoint)
        document_summary = list(enriched_documents.values())

        return {
            "total_files": len(files),
            "storage_mode": "supabase_only",
            "breakdown": {
                "original_documents": len(document_files),
                "vector_data_files": len(vector_files),
                "processed_documents": len(enriched_documents),
                "other_files": len(other_files)
            },
            "files": {
                "documents": document_files,
                "vectors": enriched_vector_files,
                "other": other_files
            },
            "processed_documents": document_summary,  # This matches your root endpoint format
            "metadata_info": {
                "documents_with_metadata": len(enriched_documents),
                "total_vector_files": len(vector_files),
                "orphaned_files": len([f for f in enriched_vector_files if f.get("file_type") == "orphaned"])
            }
        }

    except Exception as e:
        logger.error(f"Error listing storage files: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to list storage files: {str(e)}"
        )


# NEW: Standalone Supabase test endpoint
@app.get("/api/v1/test-supabase-standalone")
async def test_supabase_standalone_upload():
    """Run the standalone Supabase upload test"""
    try:
        logger.info("🧪 Running standalone Supabase upload test via API endpoint...")
        result = await test_supabase_upload_standalone()

        return {
            "status": "success" if result else "failed",
            "message": "Standalone Supabase upload test completed",
            "result": result,
            "note": "Check server logs for detailed information"
        }

    except Exception as e:
        logger.error(f"Standalone test failed: {e}")
        return {
            "status": "failed",
            "message": f"Standalone test failed: {str(e)}",
            "result": False
        }


@app.get("/api/v1/test-embedding")
async def test_embedding():
    """Test endpoint for embedding functionality"""
    try:
        test_texts = [
            "This is a test sentence for embedding.",
            "Machine learning is transforming many industries.",
            "The weather is nice today."
        ]

        # Generate embeddings
        embeddings = vector_store.embedding_model.encode(test_texts)

        # Calculate similarities
        from sklearn.metrics.pairwise import cosine_similarity
        similarities = cosine_similarity(embeddings)

        return {
            "status": "success",
            "model": EMBEDDING_MODEL_NAME,
            "embedding_dimension": vector_store.dimension,
            "test_texts": test_texts,
            "similarity_matrix": similarities.tolist(),
            "storage_mode": "supabase_only",
            "message": "Embedding service is working correctly with Supabase storage"
        }

    except Exception as e:
        logger.error(f"Embedding test failed: {e}")
        return {
            "status": "error",
            "message": f"Embedding test failed: {str(e)}"
        }


@app.get("/api/v1/test-gemini")
async def test_gemini():
    """Test Gemini API connectivity with RAG context"""
    try:
        if not rag_llm_service.model:
            return {"status": "error", "message": "Gemini model not initialized"}

        # Test with a simple RAG-style prompt
        test_context = """
        Context: Artificial Intelligence (AI) is a broad field of computer science 
        that aims to create machines capable of performing tasks that typically 
        require human intelligence.
        """

        test_question = "What is AI?"

        prompt = f"""Based on the following context, answer the question:

{test_context}

Question: {test_question}

Answer:"""

        response = rag_llm_service.model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.1,
                max_output_tokens=100,
            )
        )

        return {
            "status": "success",
            "message": "Gemini API is working with RAG",
            "model_name": rag_llm_service.model_name,
            "test_response": response.text if response and response.text else "No response text",
            "context_length": len(test_context),
            "storage_mode": "supabase_only"
        }

    except Exception as e:
        return {
            "status": "error",
            "message": f"Gemini RAG test failed: {str(e)}"
        }


@app.get("/api/v1/test-supabase")
async def test_supabase():
    """Test Supabase Storage connectivity"""
    try:
        supabase_manager = get_supabase_manager()

        # Test upload
        test_content = b"This is a test file for Supabase storage."
        test_filename = f"test_{int(time.time())}.txt"

        # Upload test file
        uploaded_path = await supabase_manager.upload_file_to_supabase(test_filename, test_content)

        # Download test file
        downloaded_content = await supabase_manager.download_file_from_supabase(uploaded_path)

        # Get public URL (now async)
        public_url = await supabase_manager.get_public_url(uploaded_path)

        # Clean up test file
        cleanup_result = await supabase_manager.delete_file_from_supabase(uploaded_path)

        return {
            "status": "success",
            "message": "Supabase Storage is working correctly",
            "test_results": {
                "upload_successful": len(uploaded_path) > 0,
                "download_successful": downloaded_content == test_content,
                "public_url_generated": len(public_url) > 0,
                "cleanup_successful": cleanup_result
            },
            "bucket": supabase_manager.bucket_name,
            "storage_mode": "supabase_only"
        }

    except Exception as e:
        logger.error(f"Supabase Storage test failed: {e}")
        return {
            "status": "error",
            "message": f"Supabase Storage test failed: {str(e)}"
        }


@app.get("/api/v1/test-vector-storage")
async def test_vector_storage():
    """Test vector storage in Supabase"""
    try:
        # Create test chunks
        test_chunks = [
            {
                "text": "This is a test chunk for vector storage testing.",
                "chunk_id": 0,
                "sentence_count": 1,
                "char_count": 49
            },
            {
                "text": "Machine learning is transforming industries rapidly.",
                "chunk_id": 1,
                "sentence_count": 1,
                "char_count": 52
            }
        ]

        # Generate test document ID
        test_doc_id = f"test_{int(time.time())}"

        # Create embeddings and save to Supabase
        chunks_created = await vector_store.create_embeddings(test_doc_id, test_chunks)

        # Test retrieval
        search_results = await vector_store.search_similar_chunks(
            test_doc_id, "machine learning", 2
        )

        # Clean up
        await vector_store.delete_document_vectors(test_doc_id)

        return {
            "status": "success",
            "message": "Vector storage in Supabase is working correctly",
            "test_results": {
                "chunks_created": chunks_created,
                "search_results_count": len(search_results),
                "cleanup_successful": True
            },
            "storage_mode": "supabase_only"
        }

    except Exception as e:
        logger.error(f"Vector storage test failed: {e}")
        return {
            "status": "error",
            "message": f"Vector storage test failed: {str(e)}"
        }


@app.middleware("http")
async def cloud_run_request_middleware(request: Request, call_next):
    """Enhanced request logging middleware optimized for Cloud Run"""
    start_time = time.time()

    # Check for graceful shutdown
    if killer.kill_now:
        return JSONResponse(
            status_code=503,
            content={"error": "Service is shutting down"}
        )

    # Log request (with truncated URL for Cloud Run logs)
    path = request.url.path
    method = request.method
    logger.info(f"Request: {method} {path}")

    try:
        # Process request
        response = await call_next(request)

        # Log response with timing
        process_time = time.time() - start_time
        logger.info(f"Response: {response.status_code} ({process_time:.3f}s)")

        # Add Cloud Run specific headers
        response.headers["X-Cloud-Run-Service"] = "rag-document-qa-api"
        response.headers["X-Response-Time"] = f"{process_time:.3f}s"

        return response

    except Exception as e:
        process_time = time.time() - start_time
        logger.error(f"Request failed: {method} {path} - {str(e)} ({process_time:.3f}s)")
        raise


# Error handlers
@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    """Enhanced HTTP exception handler"""
    logger.error(f"HTTP Exception [{request.url.path}]: {exc.status_code} - {exc.detail}")
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "error": exc.detail,
            "status_code": exc.status_code,
            "path": request.url.path,
            "storage_mode": "supabase_only"
        }
    )


@app.exception_handler(Exception)
async def general_exception_handler(request: Request, exc: Exception):
    """Enhanced general exception handler"""
    logger.error(f"Unhandled exception [{request.url.path}]: {exc}", exc_info=True)
    return JSONResponse(
        status_code=500,
        content={
            "error": "Internal server error",
            "details": str(exc) if os.getenv("DEBUG") else None,
            "path": request.url.path,
            "storage_mode": "supabase_only"
        }
    )


# Startup event
# Startup event (replace the existing startup_event function)
@app.on_event("startup")
async def startup_event():
    """Initialize services on startup with Cloud Run optimizations"""
    logger.info("🚀 Starting RAG-Powered Document QA API on Google Cloud Run")
    logger.info("📡 NO LOCAL STORAGE - All data stored in Supabase")
    logger.info(f"🧠 Embedding model: {EMBEDDING_MODEL_NAME}")
    logger.info(f"⚙️ Chunk size: {CHUNK_SIZE}, Overlap: {CHUNK_OVERLAP}")
    logger.info(f"🔍 Top-K retrieval: {TOP_K_RETRIEVAL}")
    logger.info(f"☁️ Supabase bucket: {os.getenv('SUPABASE_BUCKET', 'documents')}")
    logger.info(f"🌍 Environment: {os.getenv('ENVIRONMENT', 'development')}")
    logger.info(f"🔧 Port: {os.getenv('PORT', '8080')}")

    # Test services with shorter timeout for Cloud Run
    if rag_llm_service.model:
        logger.info(f"✅ Gemini model ready: {rag_llm_service.model_name}")
    else:
        logger.warning("⚠️ Gemini model not available")

    logger.info(f"✅ Embedding model ready: {EMBEDDING_MODEL_NAME} (dim: {vector_store.dimension})")

    # Test Supabase connection with shorter timeout for Cloud Run startup
    logger.info("🔧 Testing Supabase Storage connection...")
    try:
        # Use shorter timeout for Cloud Run
        supabase_manager = get_supabase_manager()
        files = await asyncio.wait_for(supabase_manager.list_files(), timeout=10.0)
        logger.info(f"✅ Supabase Storage connected - Found {len(files)} files")
    except asyncio.TimeoutError:
        logger.warning("⚠️ Supabase connection test timed out - continuing startup")
    except Exception as e:
        logger.warning(f"⚠️ Supabase Storage test failed: {e}")
        logger.info("API will still start - Supabase will be tested on first request")

    logger.info("✅ Cloud Run startup completed")

# Shutdown event
@app.on_event("shutdown")
async def shutdown_event():
    """Cleanup on shutdown for Cloud Run"""
    logger.info("🛑 Shutting down RAG-Powered Document QA API on Cloud Run")

    # Clear memory caches to free up resources
    if hasattr(vector_store, 'indexes'):
        vector_store.indexes.clear()
    if hasattr(vector_store, 'chunks'):
        vector_store.chunks.clear()

    logger.info("📡 All data remains safely stored in Supabase")
    logger.info("✅ Graceful shutdown completed")



if __name__ == "__main__":
    import uvicorn

    port = int(os.getenv("PORT", 8000))
    host = "0.0.0.0"

    print(f"🚀 Starting server on {host}:{port}")
    print(f"☁️ Platform: Google Cloud Run")
    print("📡 Using FULL Supabase storage - NO local files")

    # Simple uvicorn run for Cloud Run
    try:
        uvicorn.run(
            app,  # Use the app instance directly, not the string
            host=host,
            port=port,
            log_level="info",
            access_log=True,
            timeout_keep_alive=30,
            loop="auto"
        )
    except Exception as e:
        logger.error(f"❌ Server startup failed: {e}")
        sys.exit(1)
